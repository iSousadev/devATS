import io

from docx import Document


async def parse_docx(file_content: bytes) -> str:
    """Extract plain text from DOCX paragraphs and tables."""
    try:
        doc = Document(io.BytesIO(file_content))
    except Exception as exc:
        raise ValueError(f"Erro ao processar DOCX: {exc}") from exc

    parts: list[str] = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    parts.append(text)

    content = "\n".join(parts).strip()
    if not content:
        raise ValueError("DOCX nao contem texto extraivel.")

    return content
