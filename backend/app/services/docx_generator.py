import io
import re
from pathlib import Path
from datetime import datetime
from typing import Any, cast

from docx import Document as load_document
from docx.document import Document as DocxDocument
from docx.shared import Pt
from docxtpl import DocxTemplate

from app.models.schemas import ResumeData

TEMPLATES_DIR = Path(__file__).resolve().parents[1] / "templates"


def _normalize_template_id(template_id: str) -> str:
    normalized = (template_id or "").strip()
    if not normalized:
        raise ValueError("template_id e obrigatorio.")
    if normalized.lower().endswith(".docx"):
        normalized = normalized[:-5]
    return normalized


def _template_path(template_id: str) -> Path:
    normalized = _normalize_template_id(template_id)
    return TEMPLATES_DIR / f"{normalized}.docx"


def _format_date(date_str: str | None) -> str:
    if not date_str:
        return ""

    value = date_str.strip()
    if not value:
        return ""

    lowered = value.lower()
    if lowered in {"atual", "presente", "current"}:
        return "Atual"

    if re.fullmatch(r"\d{4}", value):
        return value

    if re.fullmatch(r"\d{4}-\d{2}", value):
        year, month = value.split("-")
        month_names = {
            "01": "Jan",
            "02": "Fev",
            "03": "Mar",
            "04": "Abr",
            "05": "Mai",
            "06": "Jun",
            "07": "Jul",
            "08": "Ago",
            "09": "Set",
            "10": "Out",
            "11": "Nov",
            "12": "Dez",
        }
        return f"{month_names.get(month, month)} {year}"

    return value


def _safe_filename_base(full_name: str) -> str:
    base = re.sub(r"\s+", "_", (full_name or "").strip())
    base = re.sub(r"[^A-Za-z0-9_\\-]", "", base)
    return base or "Curriculo"


def build_filename(full_name: str, yyyymmdd: str) -> str:
    return f"{_safe_filename_base(full_name)}_ATS_{yyyymmdd}.docx"


def _format_period(start_date: str | None, end_date: str | None, current: bool = False) -> str:
    start_label = _format_date(start_date) or "Inicio nao informado"
    if current:
        end_label = "Atual"
    else:
        end_label = _format_date(end_date) or "Em andamento"
    return f"{start_label} - {end_label}"


def _build_experience_block(experiences: list) -> list[dict]:
    items = []
    for exp in experiences:
        company = (exp.company or "").strip() or "Experiência não informada"
        position = (exp.position or "").strip() or "Cargo nao informado"
        period = _format_period(exp.start_date, exp.end_date, exp.current)
        location = (exp.location or "").strip()
        period_location = f"{period} | {location}" if location else period
        items.append(
            {
                "company": company,
                "position": position,
                "location": location,
                "period": period,
                "period_location": period_location,
                "achievements": [a for a in exp.achievements if a and a.strip()],
            }
        )
    return items


def _build_skills_lines(resume_data: ResumeData) -> list[str]:
    categorized = resume_data.skills.categorized or {}
    ordered = [
        ("Linguagens", categorized.get("linguagens")),
        ("Frontend", categorized.get("frontend")),
        ("Backend", categorized.get("backend")),
        ("Frameworks", categorized.get("frameworks")),
        ("Banco de Dados", categorized.get("banco_de_dados")),
        ("Ferramentas", categorized.get("ferramentas")),
        ("Práticas", categorized.get("praticas")),
    ]
    lines = [f"{label}: {value}" for label, value in ordered if value]
    if lines:
        return lines

    fallback = []
    if resume_data.skills.technical:
        fallback.append(f"Linguagens e tecnologias: {', '.join(resume_data.skills.technical)}")
    if resume_data.skills.tools:
        fallback.append(f"Ferramentas: {', '.join(resume_data.skills.tools)}")
    if resume_data.skills.soft:
        fallback.append(f"Práticas: {', '.join(resume_data.skills.soft)}")
    return fallback


def _build_context(resume_data: ResumeData) -> dict:
    headline = (resume_data.personal_info.headline or "").strip()
    if not headline:
        for exp in resume_data.experiences:
            if exp.position and exp.position.strip():
                headline = exp.position.strip()
                break

    experiences = _build_experience_block(resume_data.experiences)
    extracurricular_experiences = _build_experience_block(resume_data.extracurricular_experiences)
    education = []
    for edu in resume_data.education:
        start_label = _format_date(edu.start_date)
        end_label = _format_date(edu.end_date)
        period = ""
        current_year = datetime.now().year
        if start_label and end_label:
            period = f"{start_label} - {end_label}"
        elif end_label:
            if end_label.isdigit() and int(end_label) >= current_year:
                period = f"Em andamento (Previsão de conclusão: {end_label})"
            else:
                period = f"Conclusão: {end_label}"
        elif start_label:
            period = f"Início: {start_label}"
        education.append(
            {
                "institution": edu.institution,
                "degree": edu.degree,
                "location": edu.location or "",
                "period": period,
            }
        )

    certifications = []
    for cert in resume_data.certifications:
        date = (cert.date or "").strip()
        line = f"{cert.name} - {cert.issuer}"
        if date:
            line = f"{line} ({date})"
        certifications.append(
            {
                "line": line,
                "name": cert.name,
                "issuer": cert.issuer,
                "date": date,
                "url": str(cert.url) if cert.url else "",
            }
        )

    projects = []
    for proj in resume_data.projects:
        highlights = [h for h in proj.highlights if h and h.strip()]
        projects.append(
            {
                "name": proj.name,
                "description": proj.description,
                "technologies": ", ".join(proj.technologies),
                "technologies_dot": " · ".join(proj.technologies),
                "technologies_list": proj.technologies,
                "highlights": highlights,
                "url": str(proj.url) if proj.url else "",
            }
        )

    languages = []
    for lang in resume_data.languages:
        language = (lang.language or "").strip()
        proficiency = (lang.proficiency or "").strip()
        if not language and not proficiency:
            continue
        if not language:
            language = "Nao informado"
        if not proficiency:
            proficiency = "Nao informado"
        languages.append({"language": language, "proficiency": proficiency})

    skills_lines = _build_skills_lines(resume_data)

    return {
        "full_name": resume_data.personal_info.full_name,
        "headline": headline,
        "email": str(resume_data.personal_info.email),
        "phone": resume_data.personal_info.phone,
        "location": resume_data.personal_info.location,
        "linkedin": str(resume_data.personal_info.linkedin) if resume_data.personal_info.linkedin else "",
        "github": str(resume_data.personal_info.github) if resume_data.personal_info.github else "",
        "portfolio": str(resume_data.personal_info.portfolio) if resume_data.personal_info.portfolio else "",
        "summary": resume_data.summary or "",
        "experiences": experiences,
        "extracurricular_experiences": extracurricular_experiences,
        "education": education,
        "skills_lines": skills_lines,
        "technical_skills": ", ".join(resume_data.skills.technical),
        "tools": ", ".join(resume_data.skills.tools),
        "soft_skills": ", ".join(resume_data.skills.soft),
        "technical_skills_list": resume_data.skills.technical,
        "tools_list": resume_data.skills.tools,
        "soft_skills_list": resume_data.skills.soft,
        "certifications": certifications,
        "projects": projects,
        "languages": languages,
        "has_summary": bool(resume_data.summary),
        "has_experiences": len(experiences) > 0,
        "has_extracurricular_experiences": len(extracurricular_experiences) > 0,
        "has_education": len(resume_data.education) > 0,
        "has_skills_lines": len(skills_lines) > 0,
        "has_certifications": len(resume_data.certifications) > 0,
        "has_projects": len(resume_data.projects) > 0,
        "has_languages": len(languages) > 0,
    }


def _set_style_arial_12(doc: DocxDocument) -> None:
    normal = cast(Any, doc.styles["Normal"])
    normal.font.name = "Arial"
    normal.font.size = Pt(12)


def _remove_empty_paragraphs(doc: DocxDocument) -> None:
    for paragraph in list(doc.paragraphs):
        if paragraph.text.strip():
            continue
        element = paragraph._element
        parent = element.getparent()
        if parent is not None:
            parent.remove(element)


def _normalize_heading_text(text: str) -> str:
    replacements = {
        "Experi?ncia Profissional": "Experiência Profissional",
        "Experi?ncia Extracurricular": "Experiência Extracurricular",
        "Forma??o Acad?mica": "Formação Acadêmica",
        "Habilidades T?cnicas": "Habilidades Técnicas",
        "Experięncia Profissional": "Experiência Profissional",
        "Experięncia Extracurricular": "Experiência Extracurricular",
        "Experięncia acadęmica": "Experiência Acadêmica",
        "Forma??o Acadęmica": "Formação Acadêmica",
        "Habilidades Tęcnicas": "Habilidades Técnicas",
        "Experiencia Profissional": "Experiência Profissional",
        "Experiencia Extracurricular": "Experiência Extracurricular",
        "Formacao Academica": "Formação Acadêmica",
        "Habilidades Tecnicas": "Habilidades Técnicas",
    }
    return replacements.get(text, text)


def _postprocess_docx(docx_bytes: bytes) -> bytes:
    doc = load_document(io.BytesIO(docx_bytes))

    _set_style_arial_12(doc)

    # Normalize heading text that can come with encoding artifacts.
    for paragraph in doc.paragraphs:
        original_text = paragraph.text.strip()
        normalized_text = _normalize_heading_text(original_text)
        if normalized_text == original_text:
            continue

        had_bold = any(run.bold for run in paragraph.runs)
        paragraph.text = normalized_text
        if paragraph.runs:
            paragraph.runs[0].bold = had_bold

    heading_titles = {
        "Resumo Profissional",
        "Objetivo Profissional",
        "Experiência Profissional",
        "Experiência profissional",
        "Experiência Extracurricular",
        "Formação Acadêmica",
        "Habilidades Técnicas",
        "Tecnologias",
        "Cursos Complementares",
        "Cursos",
        "Projetos",
        "Idiomas",
    }

    first_non_empty = next((p for p in doc.paragraphs if p.text.strip()), None)
    non_empty_paragraphs = [p for p in doc.paragraphs if p.text.strip()]
    headline_paragraph = non_empty_paragraphs[1] if len(non_empty_paragraphs) > 1 else None

    # Apply font and section-aware spacing globally.
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        paragraph.paragraph_format.line_spacing = 1.15

        if not text:
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(8)
        elif paragraph is first_non_empty:
            # Nome (primeira linha)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(8)
        elif paragraph is headline_paragraph:
            # Headline (cargo)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(16)
        elif text in heading_titles:
            # Títulos de seções
            paragraph.paragraph_format.space_before = Pt(24)
            paragraph.paragraph_format.space_after = Pt(10)
        elif text.startswith(("Email:", "Telefone:", "Cidade:", "Linkedin:", "Github:", "LinkedIn:", "GitHub:", "Portfólio:")):
            # Informações de contato
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(4)
        elif " | " in text and ("-" in text or "Atual" in text or "Em andamento" in text):
            # Linha de empresa/instituição com período (ex: "Empresa X - Cargo | 2020 - 2024")
            paragraph.paragraph_format.space_before = Pt(12)
            paragraph.paragraph_format.space_after = Pt(6)
        elif text.startswith("Tecnologias usadas:") or text.startswith("Linguagens e tecnologias:") or text.startswith("Práticas:"):
            # Linha de tecnologias em projetos
            paragraph.paragraph_format.space_before = Pt(3)
            paragraph.paragraph_format.space_after = Pt(6)
        else:
            # Parágrafos normais (descrições, achievements, etc)
            paragraph.paragraph_format.space_before = Pt(3)
            paragraph.paragraph_format.space_after = Pt(8)

        for run in paragraph.runs:
            run.font.name = "Arial"
            run.font.size = Pt(12)

    _remove_empty_paragraphs(doc)

    # Name at top must be larger.
    if first_non_empty is not None:
        for run in first_non_empty.runs:
            run.font.name = "Arial"
            run.font.size = Pt(18)
            run.bold = True

    # Headline below name stays subtle.
    if len(non_empty_paragraphs) > 1:
        non_empty_paragraphs[1].paragraph_format.space_after = Pt(16)
        for run in non_empty_paragraphs[1].runs:
            run.font.name = "Arial"
            run.font.size = Pt(12)
            run.italic = True
    
    # Seções em negrito
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text in heading_titles:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(12)

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output.read()


async def generate_docx(template_id: str, resume_data: ResumeData) -> bytes:
    template_path = _template_path(template_id)

    if not template_path.exists():
        raise ValueError(f"Template '{template_id}' nao encontrado em {TEMPLATES_DIR}.")

    try:
        doc = DocxTemplate(str(template_path))
        context = _build_context(resume_data)
        doc.render(context)
    except Exception as exc:
        raise ValueError(f"Erro ao renderizar template DOCX: {exc}") from exc

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return _postprocess_docx(buffer.read())
